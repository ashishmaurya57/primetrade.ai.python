import pandas as pd
from docx import Document
import pdfkit

# Load the Excel file
file_name = "crypto_data.xlsx"
df = pd.read_excel(file_name, sheet_name="Live Data")

# Perform Analysis
top_5 = df.nlargest(5, 'market_cap')
avg_price = df['current_price'].mean()
highest_change = df.loc[df['price_change_percentage_24h'].idxmax()]
lowest_change = df.loc[df['price_change_percentage_24h'].idxmin()]

# Create a Word Report
doc = Document()
doc.add_heading("Cryptocurrency Analysis Report", 0)

doc.add_heading("1. Top 5 Cryptocurrencies by Market Capitalization", level=1)
for index, row in top_5.iterrows():
    doc.add_paragraph(f"{index+1}. {row['name']} ({row['symbol']}) - Market Cap: ${row['market_cap']:,.2f}")

doc.add_heading("2. Average Price of Top 50 Cryptocurrencies", level=1)
doc.add_paragraph(f"The average price is ${avg_price:,.2f}")

doc.add_heading("3. Highest & Lowest 24h Price Change", level=1)
doc.add_paragraph(f"Highest Gain: {highest_change['name']} ({highest_change['symbol']}) - {highest_change['price_change_percentage_24h']:.2f}%")
doc.add_paragraph(f"Biggest Drop: {lowest_change['name']} ({lowest_change['symbol']}) - {lowest_change['price_change_percentage_24h']:.2f}%")

# Save as Word file
doc.save("crypto_analysis.docx")

# Convert to PDF
pdfkit.from_file("crypto_analysis.docx", "crypto_analysis.pdf")

print("Analysis report generated: crypto_analysis.docx and crypto_analysis.pdf")
